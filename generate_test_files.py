#!/usr/bin/env python3
"""
Generate complex test documents for DIA testing
Creates realistic government documents in various formats
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

def create_government_policy_doc():
    """Create a complex government policy document (DOCX)"""
    doc = Document()
    
    # Title
    title = doc.add_heading('GOVERNMENT OF ODISHA', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Department of Rural Development', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Document header
    doc.add_paragraph('NOTIFICATION')
    doc.add_paragraph(f'No. 2456/RD/Policy/2024\t\t\t\tDate: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph()
    
    # Subject
    subject = doc.add_heading('SUBJECT: Implementation of Integrated Rural Development Scheme 2024-2029', level=3)
    
    # Main content
    doc.add_paragraph(
        'In exercise of the powers conferred under Section 23(1) of the Odisha Rural Development Act, 2019, '
        'and in supersession of all previous notifications on the subject, the State Government hereby notifies '
        'the following scheme for integrated rural development across all districts of Odisha.'
    )
    
    # Objectives section
    doc.add_heading('1. OBJECTIVES', level=2)
    objectives = [
        'To promote sustainable livelihood opportunities in rural areas through skill development and entrepreneurship',
        'To ensure universal access to basic amenities including drinking water, sanitation, and electricity',
        'To strengthen rural infrastructure including roads, bridges, and community centers',
        'To enhance agricultural productivity through modern farming techniques and access to credit',
        'To promote digital literacy and access to government services through Common Service Centers'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    # Eligibility section
    doc.add_heading('2. ELIGIBILITY CRITERIA', level=2)
    doc.add_paragraph(
        'The scheme shall be applicable to all Gram Panchayats in the state meeting the following criteria:'
    )
    eligibility = [
        'Population between 500 and 5,000 residents as per 2021 census',
        'Located in areas with less than 40% literacy rate',
        'Annual per capita income below ₹75,000',
        'Lacking basic infrastructure facilities',
        'Having at least 30% representation of women in Gram Panchayat committees'
    ]
    for item in eligibility:
        doc.add_paragraph(item, style='List Bullet')
    
    # Financial provisions
    doc.add_heading('3. FINANCIAL PROVISIONS', level=2)
    doc.add_paragraph(
        'The total budgetary allocation for this scheme is ₹1,250 crores for the financial year 2024-25, '
        'distributed as follows:'
    )
    
    # Create a table
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'S.No.'
    headers[1].text = 'Component'
    headers[2].text = 'Allocation (₹ Crores)'
    
    data = [
        ('1', 'Infrastructure Development', '450'),
        ('2', 'Skill Development & Training', '275'),
        ('3', 'Agricultural Support', '300'),
        ('4', 'Digital Infrastructure', '125'),
        ('5', 'Administrative Expenses', '100')
    ]
    
    for i, (sno, component, amount) in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = sno
        cells[1].text = component
        cells[2].text = amount
    
    # Implementation timeline
    doc.add_heading('4. IMPLEMENTATION TIMELINE', level=2)
    timeline = [
        'Phase 1 (April 2024 - September 2024): Survey and identification of eligible Gram Panchayats',
        'Phase 2 (October 2024 - March 2025): Infrastructure development and training programs',
        'Phase 3 (April 2025 - March 2026): Monitoring and mid-term evaluation',
        'Phase 4 (April 2026 - March 2029): Expansion and sustainability measures'
    ]
    for item in timeline:
        doc.add_paragraph(item, style='List Number')
    
    # Monitoring mechanism
    doc.add_heading('5. MONITORING AND EVALUATION', level=2)
    doc.add_paragraph(
        'A three-tier monitoring mechanism shall be established at District, Block, and Gram Panchayat levels. '
        'Monthly progress reports must be submitted to the State Project Management Unit. Key Performance Indicators (KPIs) include:'
    )
    kpis = [
        'Number of beneficiaries covered: Target 50,000 families',
        'Infrastructure projects completed: Minimum 80% within timeline',
        'Increase in per capita income: Expected 25% growth in 3 years',
        'Women participation: Minimum 50% in all training programs',
        'Digital literacy rate: Target 60% improvement'
    ]
    for kpi in kpis:
        doc.add_paragraph(kpi, style='List Bullet')
    
    # Penalties and violations
    doc.add_heading('6. PENALTIES FOR NON-COMPLIANCE', level=2)
    doc.add_paragraph(
        'Any Gram Panchayat or implementing agency found guilty of misappropriation of funds, '
        'falsification of records, or failure to meet minimum implementation standards shall be subject to:'
    )
    penalties = [
        'Immediate suspension of fund disbursement',
        'Recovery of misappropriated funds with 12% annual interest',
        'Blacklisting from future government schemes for a period of 5 years',
        'Legal action under relevant sections of Indian Penal Code and Prevention of Corruption Act'
    ]
    for penalty in penalties:
        doc.add_paragraph(penalty, style='List Bullet')
    
    # Grievance redressal
    doc.add_heading('7. GRIEVANCE REDRESSAL', level=2)
    doc.add_paragraph(
        'Beneficiaries may lodge complaints through:\n'
        '• State Helpline: 1800-345-6789\n'
        '• Online Portal: www.odisharuraldevelopment.gov.in\n'
        '• Email: grievance@rdodisha.gov.in\n'
        '• In-person at Block Development Office during working hours'
    )
    
    # Conclusion
    doc.add_paragraph()
    doc.add_paragraph(
        'This notification shall come into force with immediate effect and shall remain valid until March 31, 2029, '
        'unless amended or superseded by a subsequent notification.'
    )
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Signature block
    signature = doc.add_paragraph('By Order of the Governor')
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    name = doc.add_paragraph('Sd/-\nDr. Rajesh Kumar Sharma, IAS\nPrincipal Secretary\nDepartment of Rural Development\nGovernment of Odisha')
    name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Save document
    filepath = 'test_documents/government_policy_rural_development.docx'
    os.makedirs('test_documents', exist_ok=True)
    doc.save(filepath)
    print(f'✅ Created: {filepath}')
    return filepath


def create_budget_document():
    """Create a complex budget document (DOCX)"""
    doc = Document()
    
    # Title
    title = doc.add_heading('ODISHA STATE BUDGET 2024-25', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_heading('Department of Education - Annual Budget Estimates', level=2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Presented on: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph()
    
    # Executive summary
    doc.add_heading('EXECUTIVE SUMMARY', level=2)
    doc.add_paragraph(
        'The Department of Education presents the Annual Budget Estimates for FY 2024-25 with a total allocation '
        'of ₹18,500 crores, representing a 15% increase from the previous fiscal year. This budget prioritizes '
        'infrastructure development, teacher training, digital education, and inclusive education programs.'
    )
    
    # Budget highlights
    doc.add_heading('BUDGET HIGHLIGHTS', level=2)
    highlights = [
        'Total Budget: ₹18,500 crores (15% increase from FY 2023-24)',
        'Capital Expenditure: ₹5,200 crores (28% of total budget)',
        'Revenue Expenditure: ₹13,300 crores (72% of total budget)',
        'New Schools: 450 new primary and secondary schools to be established',
        'Teacher Recruitment: 12,000 new teachers across all categories',
        'Digital Infrastructure: ₹800 crores for smart classrooms and e-learning platforms'
    ]
    for item in highlights:
        doc.add_paragraph(item, style='List Bullet')
    
    # Detailed allocation
    doc.add_heading('DETAILED BUDGET ALLOCATION', level=2)
    
    table = doc.add_table(rows=11, cols=4)
    table.style = 'Medium Grid 3 Accent 1'
    
    headers = table.rows[0].cells
    headers[0].text = 'Head'
    headers[1].text = 'Sub-Category'
    headers[2].text = 'FY 2023-24 (₹ Cr)'
    headers[3].text = 'FY 2024-25 (₹ Cr)'
    
    budget_data = [
        ('Elementary Education', 'Primary Schools (Class I-V)', '4,200', '4,800'),
        ('', 'Upper Primary (Class VI-VIII)', '2,800', '3,200'),
        ('Secondary Education', 'High Schools (Class IX-X)', '2,500', '2,900'),
        ('', 'Higher Secondary (Class XI-XII)', '1,800', '2,100'),
        ('Teacher Salaries', 'Regular Teachers', '5,000', '5,500'),
        ('', 'Contractual Teachers', '800', '900'),
        ('Infrastructure', 'New Buildings & Classrooms', '1,200', '1,600'),
        ('', 'Maintenance & Repairs', '400', '500'),
        ('Digital Education', 'ICT Infrastructure', '300', '800'),
        ('Scholarships & Welfare', 'SC/ST/OBC Scholarships', '1,100', '1,200')
    ]
    
    for i, (head, subcat, prev, curr) in enumerate(budget_data, 1):
        cells = table.rows[i].cells
        cells[0].text = head
        cells[1].text = subcat
        cells[2].text = prev
        cells[3].text = curr
    
    # Comparative analysis
    doc.add_heading('COMPARATIVE ANALYSIS (Last 3 Years)', level=2)
    
    comparison = doc.add_table(rows=5, cols=4)
    comparison.style = 'Light List Accent 1'
    
    headers = comparison.rows[0].cells
    headers[0].text = 'Year'
    headers[1].text = 'Total Budget (₹ Cr)'
    headers[2].text = 'Per Student Expenditure (₹)'
    headers[3].text = 'Growth Rate (%)'
    
    comp_data = [
        ('2021-22', '14,000', '22,400', '8.5'),
        ('2022-23', '14,800', '23,650', '5.7'),
        ('2023-24', '16,100', '25,300', '8.8'),
        ('2024-25 (E)', '18,500', '28,800', '14.9')
    ]
    
    for i, (year, budget, per_student, growth) in enumerate(comp_data, 1):
        cells = comparison.rows[i].cells
        cells[0].text = year
        cells[1].text = budget
        cells[2].text = per_student
        cells[3].text = growth
    
    # New initiatives
    doc.add_heading('NEW INITIATIVES FOR FY 2024-25', level=2)
    initiatives = [
        'Launch of "Digital Odisha Schools" program in 1,000 schools with smart classrooms',
        'Establishment of 25 Model Schools in tribal and remote areas',
        'Introduction of coding and AI curriculum from Class VI onwards',
        'Free tablet distribution to 2 lakh students from economically weaker sections',
        'Teacher training program on pedagogy and technology integration (50,000 teachers)',
        'Scholarship scheme for girls pursuing STEM education (₹200 crores allocation)',
        'Mid-day meal quality improvement with nutrition monitoring system',
        'Mental health and counseling support in all secondary schools'
    ]
    for item in initiatives:
        doc.add_paragraph(item, style='List Number')
    
    # Performance metrics
    doc.add_heading('EXPECTED OUTCOMES & PERFORMANCE METRICS', level=2)
    doc.add_paragraph('The budget is designed to achieve the following targets by March 2025:')
    
    metrics = [
        'Gross Enrollment Ratio (GER) at secondary level: Increase from 76% to 82%',
        'Student-Teacher Ratio: Reduce from 32:1 to 28:1',
        'Digital literacy among students: Achieve 70% proficiency',
        'Infrastructure gap: Reduce by 40% through new construction',
        'Dropout rate: Reduce from 4.5% to 3.0% at elementary level',
        'Learning outcomes: 15% improvement in standardized test scores',
        'School attendance: Increase from 88% to 92%'
    ]
    for metric in metrics:
        doc.add_paragraph(metric, style='List Bullet')
    
    # Conclusion
    doc.add_paragraph()
    doc.add_paragraph(
        'This budget reflects the government\'s commitment to transforming education in Odisha and ensuring '
        'equitable access to quality education for all children, regardless of their socio-economic background.'
    )
    
    # Signature
    doc.add_paragraph()
    sig = doc.add_paragraph('Approved by:\nShri Arun Patel, IAS\nSecretary, Department of Education\nGovernment of Odisha')
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    filepath = 'test_documents/education_budget_2024-25.docx'
    doc.save(filepath)
    print(f'✅ Created: {filepath}')
    return filepath


def create_bilingual_notification():
    """Create a bilingual government notification (TXT)"""
    content = """
╔════════════════════════════════════════════════════════════════╗
║           GOVERNMENT OF ODISHA / ଓଡ଼ିଶା ସରକାର                    ║
║   Department of Health & Family Welfare / ସ୍ୱାସ୍ଥ୍ୟ ବିଭାଗ        ║
╚════════════════════════════════════════════════════════════════╝

NOTIFICATION / ବିଜ୍ଞପ୍ତି
No. 789/HFW/2024                    Date: December 4, 2024 / ତାରିଖ: ୪ ଡିସେମ୍ବର ୨୦୨୪

SUBJECT: Universal Health Coverage Scheme - Implementation Guidelines
ବିଷୟ: ସର୍ବଜନୀନ ସ୍ୱାସ୍ଥ୍ୟ ସୁରକ୍ଷା ଯୋଜନା - କାର୍ଯ୍ୟାନ୍ୱୟନ ନିର୍ଦ୍ଦେଶାବଳୀ

════════════════════════════════════════════════════════════════

ENGLISH VERSION:

In pursuance of the Odisha Universal Health Coverage Act, 2024, and in exercise of 
powers vested under Section 15(2) thereof, the State Government hereby notifies the 
implementation of the Universal Health Coverage Scheme across all 30 districts of Odisha.

1. SCHEME OVERVIEW
   The Universal Health Coverage Scheme (UHCS) aims to provide comprehensive healthcare
   services to all citizens of Odisha, with special focus on vulnerable populations
   including Below Poverty Line (BPL) families, senior citizens, persons with disabilities,
   and pregnant women.

2. COVERAGE DETAILS
   
   A. Beneficiaries Covered:
      • All residents of Odisha with valid Aadhaar card
      • Annual family income up to ₹5,00,000 (general category)
      • Annual family income up to ₹8,00,000 (SC/ST/OBC categories)
      • No income limit for senior citizens above 70 years
   
   B. Medical Services Included:
      • General consultation and diagnostics
      • Emergency and trauma care (24x7)
      • Surgical procedures (1,500+ procedures covered)
      • Cancer treatment and chemotherapy
      • Cardiac care and bypass surgeries
      • Dialysis and kidney transplant
      • Maternity and child care
      • Mental health services
   
   C. Financial Coverage:
      • Up to ₹5,00,000 per family per annum
      • Additional ₹5,00,000 for critical illnesses
      • No cap on emergency services
      • Cashless treatment at empaneled hospitals

3. EMPANELED HEALTHCARE FACILITIES
   
   The scheme covers 450+ hospitals across Odisha:
   • Government hospitals: 180 (including district headquarters and CHCs)
   • Private hospitals: 270 (meeting quality standards)
   • Specialized cancer centers: 12
   • Cardiac care centers: 18
   
   Complete list available at: www.odishahealthcare.gov.in

4. ENROLLMENT PROCESS
   
   Step 1: Visit nearest Common Service Center (CSC) or Aadhaar Seva Kendra
   Step 2: Submit documents:
           - Aadhaar Card (all family members)
           - Income Certificate (if applicable)
           - Ration Card
           - Bank Account details
   Step 3: Biometric verification
   Step 4: Receive Health Card within 7 working days

5. GRIEVANCE REDRESSAL
   
   Helpline Number: 104 (Toll-free, 24x7)
   Email: uhcs@odishahealth.gov.in
   Web Portal: www.odishahealthgrievance.gov.in
   
   Resolution Timeline:
   • Non-critical issues: 15 days
   • Critical issues: 7 days
   • Emergency issues: 24 hours

6. PENALTIES FOR FRAUD
   
   Any beneficiary or healthcare provider found guilty of:
   • Providing false information
   • Claiming treatment not received
   • Overcharging or demanding extra payment
   
   Shall be liable to:
   • Cancellation of health card / empanelment
   • Recovery of fraudulent claims with 18% annual interest
   • Legal action under relevant sections of IPC

This notification shall come into effect from January 1, 2025.

════════════════════════════════════════════════════════════════

ଓଡ଼ିଆ ସଂସ୍କରଣ (ODIA VERSION):

ଓଡ଼ିଶା ସର୍ବଜନୀନ ସ୍ୱାସ୍ଥ୍ୟ ସୁରକ୍ଷା ଅଧିନିୟମ, ୨୦୨୪ ଅନୁସରଣ କରି ଏବଂ ଏହାର 
ଧାରା ୧୫(୨) ଅଧୀନରେ ପ୍ରଦତ୍ତ କ୍ଷମତା ବଳରେ, ରାଜ୍ୟ ସରକାର ଓଡ଼ିଶାର ସମସ୍ତ ୩୦ 
ଜିଲ୍ଲାରେ ସର୍ବଜନୀନ ସ୍ୱାସ୍ଥ୍ୟ ସୁରକ୍ଷା ଯୋଜନାର କାର୍ଯ୍ୟକାରିତା ବିଷୟରେ ବିଜ୍ଞପ୍ତି ପ୍ରକାଶ କରୁଛନ୍ତି।

୧. ଯୋଜନା ସଂକ୍ଷିପ୍ତ ବିବରଣୀ
   ସର୍ବଜନୀନ ସ୍ୱାସ୍ଥ୍ୟ ସୁରକ୍ଷା ଯୋଜନା (UHCS) ର ଉଦ୍ଦେଶ୍ୟ ହେଉଛି ଓଡ଼ିଶାର ସମସ୍ତ ନାଗରିକଙ୍କୁ 
   ବ୍ୟାପକ ସ୍ୱାସ୍ଥ୍ୟସେବା ପ୍ରଦାନ କରିବା, ବିଶେଷ ଭାବରେ ଦାରିଦ୍ର୍ୟ ରେଖା ତଳେ (BPL) ପରିବାର, 
   ବୟସ୍କ ନାଗରିକ, ଭିନ୍ନକ୍ଷମ ବ୍ୟକ୍ତି ଏବଂ ଗର୍ଭବତୀ ମହିଳାଙ୍କ ଉପରେ ଧ୍ୟାନ ଦେବା।

୨. କଭରେଜ ବିବରଣୀ
   
   କ. ଲାଭାନ୍ୱିତ:
      • ବୈଧ ଆଧାର କାର୍ଡ ଥିବା ଓଡ଼ିଶାର ସମସ୍ତ ବାସିନ୍ଦା
      • ବାର୍ଷିକ ପାରିବାରିକ ଆୟ ₹୫,୦୦,୦୦୦ ପର୍ଯ୍ୟନ୍ତ (ସାଧାରଣ ବର୍ଗ)
      • ବାର୍ଷିକ ପାରିବାରିକ ଆୟ ₹୮,୦୦,୦୦୦ ପର୍ଯ୍ୟନ୍ତ (SC/ST/OBC)
      • ୭୦ ବର୍ଷରୁ ଅଧିକ ବୟସର ବୟସ୍କଙ୍କ ପାଇଁ କୌଣସି ଆୟ ସୀମା ନାହିଁ
   
   ଖ. ଅନ୍ତର୍ଭୁକ୍ତ ଚିକିତ୍ସା ସେବା:
      • ସାଧାରଣ ପରାମର୍ଶ ଏବଂ ନିଦାନ
      • ଜରୁରୀକାଳୀନ ଏବଂ ଆଘାତ ଚିକିତ୍ସା (୨୪x୭)
      • ଶଲ୍ୟଚିକିତ୍ସା (୧,୫୦୦+ ପ୍ରକ୍ରିୟା)
      • କର୍କଟ ଚିକିତ୍ସା ଏବଂ କେମୋଥେରାପି
      • ହୃଦୟ ଚିକିତ୍ସା ଏବଂ ବାଇପାସ ଅସ୍ତ୍ରୋପଚାର
      • ଡାୟଲିସିସ ଏବଂ କିଡନୀ ପ୍ରତିରୋପଣ
      • ମାତୃତ୍ୱ ଏବଂ ଶିଶୁ ଯତ୍ନ
      • ମାନସିକ ସ୍ୱାସ୍ଥ୍ୟ ସେବା
   
   ଗ. ଆର୍ଥିକ କଭରେଜ:
      • ପ୍ରତି ପରିବାର ପ୍ରତି ବର୍ଷ ₹୫,୦୦,୦୦୦ ପର୍ଯ୍ୟନ୍ତ
      • ଗୁରୁତର ରୋଗ ପାଇଁ ଅତିରିକ୍ତ ₹୫,୦୦,୦୦୦
      • ଜରୁରୀକାଳୀନ ସେବା ଉପରେ କୌଣସି ସୀମା ନାହିଁ
      • ତାଲିକାଭୁକ୍ତ ହସ୍ପିଟାଲରେ କ୍ୟାସଲେସ ଚିକିତ୍ସା

୩. ସ୍ୱାସ୍ଥ୍ୟସେବା ସୁବିଧା
   ଯୋଜନାରେ ଓଡ଼ିଶା ମଧ୍ୟରେ ୪୫୦+ ହସ୍ପିଟାଲ ଅନ୍ତର୍ଭୁକ୍ତ:
   • ସରକାରୀ ହସ୍ପିଟାଲ: ୧୮୦
   • ବ୍ୟକ୍ତିଗତ ହସ୍ପିଟାଲ: ୨୭୦
   • ବିଶେଷଜ୍ଞ କର୍କଟ କେନ୍ଦ୍ର: ୧୨
   • ହୃଦୟ ଚିକିତ୍ସା କେନ୍ଦ୍ର: ୧୮

୪. ନାମଲେଖା ପ୍ରକ୍ରିୟା
   ପଦକ୍ଷେପ ୧: ନିକଟତମ କମନ୍ ସର୍ଭିସ ସେଣ୍ଟର (CSC) କୁ ଯାଆନ୍ତୁ
   ପଦକ୍ଷେପ ୨: ଦଲିଲ ଦାଖଲ କରନ୍ତୁ:
            - ଆଧାର କାର୍ଡ (ସମସ୍ତ ପରିବାର ସଦସ୍ୟ)
            - ଆୟ ପ୍ରମାଣପତ୍ର
            - ରେସନ କାର୍ଡ
            - ବ୍ୟାଙ୍କ ଖାତା ବିବରଣୀ
   ପଦକ୍ଷେପ ୩: ବାୟୋମେଟ୍ରିକ ଯାଞ୍ଚ
   ପଦକ୍ଷେପ ୪: ୭ କାର୍ଯ୍ୟଦିବସ ମଧ୍ୟରେ ସ୍ୱାସ୍ଥ୍ୟ କାର୍ଡ ପାଆନ୍ତୁ

୫. ଅଭିଯୋଗ ନିବାରଣ
   ହେଲ୍ପଲାଇନ୍: ୧୦୪ (ଟୋଲ-ଫ୍ରି, ୨୪x୭)
   ଇମେଲ: uhcs@odishahealth.gov.in
   ୱେବ୍ ପୋର୍ଟାଲ: www.odishahealthgrievance.gov.in

ଏହି ବିଜ୍ଞପ୍ତି ୧ ଜାନୁଆରୀ, ୨୦୨୫ ଠାରୁ ଲାଗୁ ହେବ।

════════════════════════════════════════════════════════════════

For queries, contact:
ପ୍ରଶ୍ନ ପାଇଁ ଯୋଗାଯୋଗ କରନ୍ତୁ:

State Health Mission Office / ରାଜ୍ୟ ସ୍ୱାସ୍ଥ୍ୟ ମିଶନ କାର୍ଯ୍ୟାଳୟ
Unit-5, Bhubaneswar, Odisha - 751001
Phone: 0674-2536363
Email: shm.odisha@gov.in

By Order / ଆଦେଶକ୍ରମେ,
Dr. Sanjay Mishra, IAS
Commissioner-cum-Secretary / କମିଶନର-ସହ-ସଚିବ
Health & Family Welfare Department / ସ୍ୱାସ୍ଥ୍ୟ ଏବଂ ପରିବାର କଲ୍ୟାଣ ବିଭାଗ
Government of Odisha / ଓଡ଼ିଶା ସରକାର
"""
    
    filepath = 'test_documents/health_notification_bilingual.txt'
    os.makedirs('test_documents', exist_ok=True)
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f'✅ Created: {filepath}')
    return filepath


def create_procurement_order():
    """Create a government procurement order (TXT)"""
    content = """
╔════════════════════════════════════════════════════════════════════╗
║                    GOVERNMENT OF ODISHA                            ║
║           Public Works Department (Roads & Buildings)              ║
║                 Tender & Procurement Division                      ║
╚════════════════════════════════════════════════════════════════════╝

TENDER NOTICE
Tender No: PWD/RB/2024/TN-45678
Issue Date: December 4, 2024
Submission Deadline: January 15, 2025, 3:00 PM IST

════════════════════════════════════════════════════════════════════

PROJECT: Construction of State Highway Expansion - Phases 3 & 4
         (Bhubaneswar to Puri via Konark - 120 KM Corridor)

════════════════════════════════════════════════════════════════════

1. TENDER OVERVIEW

The Public Works Department, Government of Odisha, invites sealed tenders from
eligible contractors for the design, construction, and maintenance of a four-lane
highway connecting Bhubaneswar to Puri via Konark tourist route.

Estimated Project Cost: ₹3,450 Crores
Contract Duration: 36 months (including monsoon delays)
Defect Liability Period: 60 months post-completion

════════════════════════════════════════════════════════════════════

2. SCOPE OF WORK

Phase 3 (Bhubaneswar to Konark - 65 KM):
• Widening of existing 2-lane road to 4-lane divided highway
• Construction of 8 major bridges and 24 minor bridges/culverts
• Installation of toll collection systems (4 toll plazas)
• Service roads along the entire stretch
• Street lighting with solar-powered LED systems
• Roadside amenities: Rest areas, fuel stations, emergency call boxes
• Green belt development along both sides (50m width)

Phase 4 (Konark to Puri - 55 KM):
• New 4-lane highway construction (greenfield project)
• Bypass around Konark temple area (heritage zone)
• Construction of 1 major flyover and 3 underpasses
• Dedicated pedestrian and cycling tracks
• Tourist information centers at 3 locations
• Parking facilities near beach areas
• Drainage and flood management systems

Additional Requirements:
• Quality control laboratory setup at project site
• Traffic management during construction
• Environmental mitigation measures
• Relocation assistance for affected parties (if any)
• Safety measures and signage as per IRC standards

════════════════════════════════════════════════════════════════════

3. TECHNICAL SPECIFICATIONS

Road Design:
• Pavement: 250mm thick bituminous surface
• Base course: WBM (Water Bound Macadam) - 150mm
• Sub-base: Granular sub-base - 200mm
• Median width: 5m with crash barriers
• Carriageway width: 7m per lane
• Shoulder width: 2.5m (paved)
• Design speed: 100 km/h
• Design life: 30 years

Bridge Specifications:
• Load capacity: IRC Class-A or AA loading
• Seismic zone considerations (Zone III)
• Foundation depth as per soil investigation reports
• Use of high-quality concrete (M40 grade minimum)
• Corrosion-resistant steel reinforcement

Materials Quality Standards:
• All materials must conform to IS codes and MORT&H specifications
• Aggregate: Conforming to IS 2386
• Cement: OPC 53 grade conforming to IS 12269
• Bitumen: VG-30 grade conforming to IS 73
• Steel: Fe 500D conforming to IS 1786

════════════════════════════════════════════════════════════════════

4. ELIGIBILITY CRITERIA

Financial Capability:
• Minimum average annual turnover: ₹500 crores (last 3 financial years)
• Net worth: Minimum ₹200 crores as of March 31, 2024
• Working capital: Minimum ₹150 crores
• Credit rating: Minimum BBB+ from CRISIL or equivalent

Technical Capability:
• Experience in highway construction: Minimum 15 years
• Completed projects:
  - At least 3 highway projects of 50+ KM length in last 7 years
  - At least 1 project worth ₹1,000+ crores
  - Experience in bridge construction (span > 50m)
• In-house equipment: As per detailed list in tender document
• Quality certifications: ISO 9001:2015, ISO 14001:2015

Personnel Requirements:
• Project Manager: Degree in Civil Engineering with 20+ years experience
• Quality Control Manager: Degree with 15+ years experience
• Safety Officer: Certified professional with 10+ years experience
• Environmental Officer: As per environmental clearance norms
• Minimum workforce capacity: 500 skilled workers

Legal Requirements:
• Valid contractor license from PWD Odisha (Class-I)
• GST registration
• PAN and TAN
• EPF and ESI registration
• Valid pollution clearance certificate
• No ongoing disputes or blacklisting

════════════════════════════════════════════════════════════════════

5. COST BREAKDOWN & PAYMENT TERMS

Estimated Cost Distribution:

Component                                  Amount (₹ Crores)    % of Total
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Earthwork & Site Preparation               420                  12.2%
Pavement & Surfacing                       980                  28.4%
Bridge & Culvert Construction              785                  22.8%
Drainage Systems                           245                  7.1%
Traffic Management & Signage               180                  5.2%
Toll Plaza Infrastructure                  325                  9.4%
Electrical & Street Lighting               215                  6.2%
Environmental & Safety Measures            125                  3.6%
Contingency & Price Escalation             175                  5.1%
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
TOTAL                                      3,450                100%

Payment Schedule:
• Mobilization advance: 10% of contract value (after bank guarantee)
• Running bills: Monthly based on work progress
• Retention money: 10% of each bill (released after defect liability period)
• Final payment: 95% on provisional completion, 5% on final completion

Price Variation:
• Applicable as per MORT&H formula for cement, steel, bitumen, and POL
• Base date for price calculation: 28 days before tender submission deadline

════════════════════════════════════════════════════════════════════

6. SUBMISSION REQUIREMENTS

Earnest Money Deposit (EMD):
• Amount: ₹35 crores (1% of estimated cost)
• Form: Bank guarantee from scheduled commercial bank
• Validity: 180 days from tender submission date

Technical Bid Documents:
1. Company profile and registration certificates
2. Financial statements for last 3 years (audited)
3. Experience certificates for similar projects
4. Details of ongoing projects
5. Equipment list with ownership proof
6. Key personnel CVs and experience certificates
7. Methodology and work plan
8. Quality assurance plan
9. Environmental management plan
10. Safety plan

Financial Bid:
• To be submitted in separate sealed envelope
• Bill of Quantities (BOQ) as per tender document
• Detailed rate analysis for major items
• Breakdown of overhead charges

Bank Guarantees Required:
• EMD: ₹35 crores
• Performance guarantee: 10% of contract value (after award)
• Mobilization advance guarantee: 10% of contract value
• Retention money guarantee: Optional (in lieu of retention)

════════════════════════════════════════════════════════════════════

7. EVALUATION CRITERIA

Stage 1 - Preliminary Screening:
• Eligibility criteria verification
• Completeness of documents
• EMD verification

Stage 2 - Technical Evaluation (70 marks):
• Past experience (20 marks)
• Financial capability (15 marks)
• Equipment and resources (15 marks)
• Methodology and work plan (10 marks)
• Quality and safety plan (10 marks)

Minimum qualifying score: 49 marks (70%)

Stage 3 - Financial Evaluation (30 marks):
• L1 bidder: 30 marks
• Others: Marks proportional to L1 bid

Final Selection:
• Combined score of technical (70%) + financial (30%)
• Highest scorer wins the contract
• Arithmetic errors correction as per tender conditions

════════════════════════════════════════════════════════════════════

8. KEY DATES & DEADLINES

Pre-bid Meeting:           December 18, 2024, 11:00 AM
                          (PWD Head Office, Bhubaneswar)

Site Visit:               December 20-22, 2024
                          (Registration required)

Last Date for Queries:    December 28, 2024, 5:00 PM

Clarification Response:   January 3, 2025

Tender Submission:        January 15, 2025, 3:00 PM
                          (Online through e-procurement portal)

Technical Bid Opening:    January 16, 2025, 11:00 AM

Financial Bid Opening:    To be notified after technical evaluation

Award of Contract:        Within 60 days of tender opening

Project Commencement:     Within 30 days of award

════════════════════════════════════════════════════════════════════

9. SPECIAL CONDITIONS

Environmental Compliance:
• Project has received environmental clearance (No. EC/2024/ORD/8756)
• Contractor must comply with all environmental norms
• Regular environmental monitoring reports required
• Tree plantation: 5 trees for every tree cut
• Noise and dust control measures mandatory

Social Impact:
• Employment preference to local youth (minimum 30%)
• Skill development training for 500 local workers
• CSR activities in project-affected areas
• Grievance redressal mechanism for local community

Penalties:
• Delay in completion: 0.05% of contract value per day (max 10%)
• Quality defects: As per defect liability clause
• Safety violations: ₹1 lakh per incident
• Environmental violations: ₹5 lakhs per incident

Dispute Resolution:
• Amicable settlement through project management committee
• Arbitration as per Arbitration and Conciliation Act, 1996
• Jurisdiction: Courts in Bhubaneswar, Odisha

════════════════════════════════════════════════════════════════════

10. CONTACT INFORMATION

Tender Documents:
Available for download at: https://tenderodisha.gov.in
Cost of tender document: ₹50,000 (non-refundable)
Document can be purchased online through e-payment

Technical Queries:
Chief Engineer (Roads)
Public Works Department
Odisha Secretariat, Bhubaneswar - 751001
Email: ce.roads@pwdodisha.gov.in
Phone: 0674-2536789

Procurement Queries:
Superintending Engineer (Tender Cell)
Email: tender.pwd@odisha.gov.in
Phone: 0674-2536790

Helpdesk:
Available Monday to Friday, 10:00 AM - 5:00 PM
Helpline: 1800-345-7890 (Toll-free)

════════════════════════════════════════════════════════════════════

GENERAL CONDITIONS:
1. The department reserves the right to accept or reject any or all tenders
   without assigning any reason.
2. Conditional bids will be summarily rejected.
3. Canvassing in any form will result in disqualification.
4. All disputes are subject to Bhubaneswar jurisdiction only.
5. Tender documents are not transferable.

════════════════════════════════════════════════════════════════════

Issued by:
Engineer-in-Chief
Public Works Department (Roads & Buildings)
Government of Odisha

Date: December 4, 2024
Place: Bhubaneswar

[Digitally Signed]
Seal: Government of Odisha, PWD
"""
    
    filepath = 'test_documents/procurement_tender_highway.txt'
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f'✅ Created: {filepath}')
    return filepath


def main():
    """Generate all test documents"""
    print("\n" + "="*70)
    print("  GENERATING TEST DOCUMENTS FOR DIA")
    print("="*70 + "\n")
    
    files = []
    
    try:
        # Generate DOCX files
        print("📄 Generating DOCX documents...")
        files.append(create_government_policy_doc())
        files.append(create_budget_document())
        
        # Generate TXT files
        print("\n📝 Generating TXT documents...")
        files.append(create_bilingual_notification())
        files.append(create_procurement_order())
        
        print("\n" + "="*70)
        print("✅ SUCCESSFULLY GENERATED ALL TEST DOCUMENTS")
        print("="*70)
        print(f"\nTotal files created: {len(files)}")
        print("\nFiles location: test_documents/")
        print("\nGenerated files:")
        for i, file in enumerate(files, 1):
            print(f"  {i}. {file}")
        
        print("\n" + "="*70)
        print("🧪 SUGGESTED TESTS:")
        print("="*70)
        print("\n1. SUMMARIZE TEST:")
        print("   - Upload: government_policy_rural_development.docx")
        print("   - Task: Summarize")
        print("   - Language: English")
        
        print("\n2. EXTRACT TEST:")
        print("   - Upload: education_budget_2024-25.docx")
        print("   - Task: Extract")
        print("   - Query: Extract all budget allocations and their amounts")
        
        print("\n3. BILINGUAL TEST:")
        print("   - Upload: health_notification_bilingual.txt")
        print("   - Task: Summarize")
        print("   - Language: Bilingual")
        
        print("\n4. Q&A TEST:")
        print("   - Upload: procurement_tender_highway.txt")
        print("   - Task: Q&A")
        print("   - Query: What are the eligibility criteria for bidders?")
        
        print("\n5. COMPARE TEST:")
        print("   - Upload: government_policy_rural_development.docx")
        print("   - Upload: education_budget_2024-25.docx")
        print("   - Task: Compare")
        print("   - Language: English")
        
        print("\n" + "="*70)
        
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
